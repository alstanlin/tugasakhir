import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_qna(doc, category, questions):
    add_heading(doc, category, 2)
    for i, q in enumerate(questions):
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"Q{i+1}: {q['question']}")
        run_q.bold = True
        run_q.font.color.rgb = RGBColor(192, 0, 0)
        
        doc.add_paragraph(q['answer'])
        
        p_tips = doc.add_paragraph()
        run_tips = p_tips.add_run("💡 Tips Menjawab: ")
        run_tips.bold = True
        run_tips.font.color.rgb = RGBColor(0, 128, 0)
        p_tips.add_run(q['tips'])
        
        p_mistakes = doc.add_paragraph()
        run_mistakes = p_mistakes.add_run("⚠️ Hindari: ")
        run_mistakes.bold = True
        run_mistakes.font.color.rgb = RGBColor(255, 140, 0)
        p_mistakes.add_run(q['mistakes'])
        
        doc.add_paragraph("_" * 50)

doc = docx.Document()

# Title
title = doc.add_heading('Q&A Persiapan Sidang Tugas Akhir', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Nama: Alfonso Liguori Stanlin Dyas")
doc.add_paragraph("NIM: 1105220040")
doc.add_paragraph("Judul: Analisis Kinerja Wearable Elektrokardiogram Menggunakan Elektroda Nanokomposit Berbasis Convolutional Neural Network untuk Mendeteksi Aritmia")
doc.add_paragraph("\nDokumen ini berisi kompilasi pertanyaan fundamental dan spesifik berdasarkan data, metodologi, dan hasil aktual dari pengerjaan Tugas Akhir.")
doc.add_page_break()

data = {
    "A. PERTANYAAN FUNDAMENTAL (Teori Dasar)": [
        {
            "question": "Apa itu SNR dan bagaimana cara menginterpretasikan nilai SNR 1.80 dB pada elektroda nanokomposit Anda?",
            "answer": "SNR (Signal-to-Noise Ratio) adalah rasio perbandingan antara kekuatan sinyal (informasi) dengan kekuatan noise (gangguan). Pada pengukuran logaritmik (dB), rumus yang digunakan adalah 10 * log10(P_signal / P_noise). Nilai SNR 1.80 dB pada elektroda nanokomposit berarti sinyal EKG hanya sekitar 1.5 kali lebih kuat dari noise. Ini adalah tingkat noise yang sangat tinggi (hampir menutupi sinyal asli), yang disebabkan oleh impedansi kontak yang tinggi pada dry electrode (nanokomposit) tanpa penggunaan gel konduktif.",
            "tips": "Jelaskan dengan lugas bahwa 1.80 dB itu sangat rendah. Kaitkan dengan realita dry electrode yang memang rentan terhadap motion artifact dan noise impedansi.",
            "mistakes": "Jangan mengatakan 1.80 dB itu nilai yang bagus. Nilai standar EKG klinis biasanya di atas 20 dB."
        },
        {
            "question": "Mengapa Anda menggunakan frekuensi sampling (FS) 250 Hz? Hubungkan dengan Teorema Nyquist!",
            "answer": "Berdasarkan Teorema Nyquist-Shannon, frekuensi sampling harus minimal dua kali lipat dari frekuensi tertinggi sinyal yang ingin diamati (Fs ≥ 2 * Fmax). Mayoritas informasi diagnostik EKG (PQRST) berada pada rentang frekuensi 0.5 Hz hingga 40 Hz (bahkan klinis sampai 100 Hz). Dengan FS = 250 Hz, kita bisa merekonstruksi frekuensi sinyal hingga 125 Hz tanpa efek aliasing, yang sudah sangat mencukupi untuk menangkap seluruh morfologi EKG dengan detail.",
            "tips": "Sebutkan Fs = 250 Hz berarti Nyquist frequency-nya 125 Hz. Frekuensi sinyal EKG kita (setelah bandpass) dibatasi sampai 40 Hz.",
            "mistakes": "Lupa menyebutkan batas frekuensi maksimal sinyal EKG."
        },
        {
            "question": "Mengapa menggunakan Conv1D? Apa bedanya dengan Conv2D?",
            "answer": "Conv1D digunakan untuk data sekuensial satu dimensi seperti sinyal time-series (EKG) atau teks, di mana filter (kernel) hanya bergeser searah sumbu waktu. Sedangkan Conv2D digunakan untuk data spasial dua dimensi seperti gambar (geser horizontal dan vertikal). EKG adalah sinyal satu dimensi (amplitudo terhadap waktu), sehingga Conv1D adalah pilihan paling natural dan efisien untuk mengekstrak pola lokal seperti bentuk gelombang QRS, P, dan T secara otomatis.",
            "tips": "Tekankan pada bentuk data EKG yang merupakan deret waktu (time-series).",
            "mistakes": "Menyebutkan bahwa gambar EKG di-crop (kecuali Anda memang mengubah sinyal jadi gambar spektrogram, padahal di TA ini pakai sinyal 1D raw)."
        }
    ],
    "B. BAB I - PENDAHULUAN & BAB II - TINJAUAN PUSTAKA": [
        {
            "question": "Di judul tertulis 'Mendeteksi Aritmia', namun data yang Anda gunakan adalah data simulasi aritmia oleh subjek sehat. Apakah ini tidak menyalahi klaim deteksi?",
            "answer": "Itu adalah keterbatasan penelitian ini. Pengambilan data dilakukan pada orang sehat yang mensimulasikan ritme jantung tidak beraturan (tapping). Secara klinis, ini bukan aritmia patologis (seperti AFib sejati), melainkan 'irregular heartbeat' yang disimulasikan secara mekanik. Saya menyadari hal ini menyebabkan domain shift ketika model dicoba ke data pasien asli (MIT-BIH). Judul ini lebih berfokus pada pengujian sistem dari segi instrumentasi dan algoritma untuk membedakan ritme reguler vs irreguler.",
            "tips": "Jujur tentang limitasi data primer. Jangan mengklaim sistem sudah siap klinis.",
            "mistakes": "Berdalih bahwa data yang disimulasikan sama persis secara biologis dengan AFib pasien."
        },
        {
            "question": "Mengapa memilih metode wavelet (DWT) db5 untuk denoising?",
            "answer": "Wavelet Daubechies 5 (db5) dipilih karena memiliki bentuk fungsi basis (mother wavelet) yang asimetris, yang sangat menyerupai morfologi sinyal EKG, khususnya bentuk kompleks QRS. Hal ini memungkinkan dekomposisi energi yang lebih optimal dibandingkan wavelet lain (seperti Haar atau Symlet). Selain itu, dengan mendekomposisi hingga level 5, kita dapat memisahkan pita frekuensi noise (seperti baseline wander di frekuensi sangat rendah dan noise frekuensi tinggi) untuk kemudian di-threshold dengan efektif tanpa merusak morfologi R-peak.",
            "tips": "Sebutkan keyword 'kemiripan bentuk dengan kompleks QRS' (morphological similarity).",
            "mistakes": "Mengatakan db5 adalah yang terbaik untuk semua sinyal (hanya cocok karena mirip EKG)."
        }
    ],
    "C. BAB III - METODOLOGI": [
        {
            "question": "Pada DWT Anda mengenolkan koefisien cA5 (coeffs[0] = 0). Apa tujuan dari langkah ini?",
            "answer": "Pada dekomposisi DWT level 5 (FS = 250 Hz), koefisien aproksimasi level 5 (cA5) mewakili rentang frekuensi yang sangat rendah (sekitar 0 - 3.9 Hz). Komponen baseline wander (osilasi dasar pernapasan) mayoritas berada di frekuensi bawah 0.5 Hz. Dengan menolkan cA5, secara efektif kita menghilangkan baseline wander pada sinyal, sehingga sinyal EKG menjadi lurus (isoelectric line stabil) tanpa merusak gelombang cepat seperti kompleks QRS.",
            "tips": "Hubungkan cA5 dengan frekuensi sangat rendah dan baseline wander.",
            "mistakes": "Mengatakan cA5 menghilangkan motion artifact (motion artifact ada di berbagai pita frekuensi)."
        },
        {
            "question": "Bagaimana proses labeling data primer Anda dilakukan? Apakah valid?",
            "answer": "Labeling dilakukan secara manual menggunakan GUI script (labeling_nano.py) yang saya kembangkan, di mana setiap window 10 detik dilabeli sebagai Normal atau AFib (simulasi). Secara akademis, ini memiliki risiko bias karena tidak divalidasi oleh dokter kardiologi spesialis (Gold Standard). Oleh karena itu, label ini didasarkan pada definisi teknis (keteraturan interval R-R), bukan diagnosis medis. Ini adalah kelemahan yang saya tulis di saran untuk penelitian lanjutan.",
            "tips": "Akui kelemahan secara terbuka. Dosen lebih menghargai mahasiswa yang sadar akan batasan risetnya.",
            "mistakes": "Bersikeras bahwa label mahasiswa sudah pasti 100% benar secara medis."
        },
        {
            "question": "Mengapa menggunakan SMOTE-Tomek? Bukankah itu merusak data asli?",
            "answer": "Data primer saya sangat tidak seimbang (699 Normal, 124 AFib). Jika dilatih langsung, model akan bias memprediksi Normal dan mengabaikan kelas minoritas (Sensitivitas jatuh). SMOTE-Tomek digunakan hanya pada data TRAINING. SMOTE men-sintesis data minoritas baru menggunakan interpolasi jarak K-Nearest Neighbors, sedangkan Tomek Links membersihkan batas antar kelas dengan menghapus sampel mayoritas yang terlalu dekat dengan minoritas. Ini menjaga batasan keputusan (decision boundary) tetap bersih.",
            "tips": "Sangat krusial: Tekankan bahwa SMOTE HANYA diterapkan pada data training. Data testing 100% asli.",
            "mistakes": "Menyebutkan bahwa SMOTE diterapkan sebelum splitting data (ini disebut data leakage, kesalahan fatal!)."
        }
    ],
    "D. BAB IV - HASIL & PEMBAHASAN": [
        {
            "question": "Akurasi training Anda hampir 100%, tapi akurasi validasi berhenti di 93% dan loss memburuk. Apakah model Anda overfitting?",
            "answer": "Ya, model saya mengalami overfitting yang signifikan pada dataset primer. Hal ini terlihat dari divergensi antara training loss yang terus turun (sampai 2e-5) sementara validation loss mulai naik (plateau di 0.22 sejak epoch 3). Overfitting ini disebabkan oleh dua hal utama: pertama, jumlah dataset yang sangat kecil (hanya 823 window), dan kedua, kualitas sinyal dengan SNR sangat rendah (1.80 dB) yang memaksa model 'menghafal' noise pada data training alih-alih mempelajari pola umum EKG yang bersih.",
            "tips": "Jangan takut mengakui overfitting. Analisis mengapa itu terjadi (data dikit, noise tinggi).",
            "mistakes": "Membantah overfitting hanya karena akurasi testing 91%. (Validation loss yang naik adalah bukti valid overfitting)."
        },
        {
            "question": "Precision 100% tapi Sensitivity (Recall) hanya 68.42%. Apa artinya ini untuk alat deteksi medis?",
            "answer": "Precision 100% berarti alat saya tidak pernah salah menuduh orang sehat sebagai AFib (False Positive = 0). Namun, Sensitivity 68.42% berarti alat saya gagal mendeteksi sekitar 31% kasus AFib (False Negative = 6). Dalam konteks medis, ini berbahaya karena pasien AFib bisa pulang mengira mereka sehat. Sistem skrining medis umumnya membutuhkan Sensitivity yang tinggi agar tidak melewatkan pasien sakit, bahkan jika harus mengorbankan sedikit Precision.",
            "tips": "Pahami tradeoff Precision vs Recall di ranah medis (Recall lebih penting untuk deteksi penyakit).",
            "mistakes": "Fokus pada angka akurasi 91% tanpa membahas Recall yang rendah."
        },
        {
            "question": "Mengapa performa pada data primer sangat jauh berbeda dibandingkan dengan benchmark MIT-BIH?",
            "answer": "Perbedaan performa disebabkan oleh kualitas data. Dataset MIT-BIH (yang mencapai akurasi 98.7% dan sensitivitas 97.9%) direkam dengan elektroda klinis (wet electrode) pada rasio SNR yang jauh lebih baik dan telah dianotasi oleh kardiolog profesional. Sedangkan data primer menggunakan elektroda nanokomposit tipe dry electrode dengan SNR yang sangat rendah (1.80 dB) akibat tingginya impedansi kontak, serta pola aritmia yang hanya simulasi. Arsitektur CAT-Net terbukti sangat andal (teruji di MIT-BIH), namun model terhambat oleh kualitas input (GIGO: Garbage In, Garbage Out).",
            "tips": "Salurkan 'kesalahan' bukan pada model AI, tetapi pada keterbatasan akuisisi hardware (elektroda) dan jumlah data.",
            "mistakes": "Menyalahkan algoritma AI tanpa merujuk keberhasilannya di dataset MIT-BIH."
        }
    ],
    "E. BAB V - KESIMPULAN & SARAN": [
        {
            "question": "Apakah sistem Anda sudah layak disebut berhasil?",
            "answer": "Secara instrumentasi (kemampuan mengakuisisi) dan konseptual arsitektur (terbukti pada MIT-BIH), sistem ini berhasil. Namun, secara klinis untuk penggunaan mandiri pada data nanokomposit, sistem ini belum sepenuhnya layak. Kelemahan pada SNR yang rendah (1.80 dB) dan sensitivitas yang terbatas (68%) menunjukkan bahwa sistem memerlukan optimalisasi pada sisi material elektroda (penurunan impedansi kontak) sebelum algoritma AI dapat bekerja secara maksimal.",
            "tips": "Jawab dengan objektif. Pisahkan keberhasilan sistem (secara arsitektur) dengan batasan klinisnya.",
            "mistakes": "Klaim over-promising bahwa alat ini siap diproduksi massal atau dipakai RS."
        }
    ]
}

for cat, questions in data.items():
    add_qna(doc, cat, questions)

doc.save('QnA_Sidang_TA_Alfonso.docx')
print("DOCX successfully generated.")
